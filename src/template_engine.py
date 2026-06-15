"""Generate professional DOCX proposals"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from pathlib import Path
from .models import ExtractedData

class ProposalGenerator:
    def __init__(self):
        """Initialize the generator"""
        pass
    
    def generate(self, data: ExtractedData, output_path: str) -> str:
        """Generate a proposal document"""
        
        # Create a new document
        doc = Document()
        
        # Add Title
        title = doc.add_heading("BUSINESS PROPOSAL", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Empty line
        
        # Add date
        date_para = doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()
        
        # Client section
        doc.add_heading("Prepared For", level=1)
        doc.add_paragraph(data.client_name or "[Client Name]")
        
        doc.add_paragraph()
        
        # Event Overview
        doc.add_heading("Event Overview", level=1)
        
        # Create a table for event details
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Detail"
        hdr_cells[1].text = "Information"
        
        # Add event details
        details = [
            ("Event Name", data.event_name or "To be confirmed"),
            ("Event Date", data.event_date or "To be confirmed"),
            ("Venue", data.venue or "To be confirmed"),
            ("Expected Guests", str(data.guest_count) if data.guest_count else "To be confirmed"),
        ]
        
        for detail, value in details:
            row = table.add_row()
            row.cells[0].text = detail
            row.cells[1].text = value
        
        doc.add_paragraph()
        
        # Executive Summary
        doc.add_heading("Executive Summary", level=1)
        summary = f"""
        This proposal outlines the services and arrangements for {data.event_name or 'your upcoming event'}.
        Based on our discussion, we have prepared the following plan to ensure a successful event.
        """
        doc.add_paragraph(summary.strip())
        
        doc.add_paragraph()
        
        # Services Section
        doc.add_heading("Proposed Services", level=1)
        
        if data.services_requested:
            for service in data.services_requested:
                doc.add_paragraph(f"• {service}", style='List Bullet')
        else:
            doc.add_paragraph("Services to be confirmed based on final requirements.")
        
        doc.add_paragraph()
        
        # Special Requirements
        if data.special_requirements:
            doc.add_heading("Special Requirements", level=1)
            for req in data.special_requirements:
                doc.add_paragraph(f"• {req}", style='List Bullet')
            doc.add_paragraph()
        
        # Pricing Section
        doc.add_heading("Investment", level=1)
        
        if data.budget:
            doc.add_paragraph(f"Client Budget Indication: {data.budget}")
            doc.add_paragraph()
        
        pricing_text = """
        Pricing Structure:
        
        • 50% deposit to confirm booking
        • 25% due 30 days prior to event
        • 25% due on event day
        
        *All pricing excludes applicable taxes.
        """
        doc.add_paragraph(pricing_text)
        
        doc.add_paragraph()
        
        # Terms
        doc.add_heading("Terms & Conditions", level=1)
        terms = """
        1. This proposal is valid for 30 days from the date of issue.
        2. Changes to scope may affect pricing.
        3. Cancellation policy applies as per signed agreement.
        4. All prices are in USD unless specified otherwise.
        """
        doc.add_paragraph(terms)
        
        doc.add_paragraph()
        
        # Signature
        doc.add_heading("Acceptance", level=1)
        sig_table = doc.add_table(rows=3, cols=2)
        sig_table.autofit = False
        sig_table.columns[0].width = Inches(3)
        sig_table.columns[1].width = Inches(3)
        
        sig_table.cell(0, 0).text = "Accepted by:"
        sig_table.cell(0, 1).text = "_________________________"
        sig_table.cell(1, 0).text = "Title:"
        sig_table.cell(1, 1).text = "_________________________"
        sig_table.cell(2, 0).text = "Date:"
        sig_table.cell(2, 1).text = "_________________________"
        
        # Save the document
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        
        print(f"✅ Proposal saved to: {output_path}")
        return str(output_path)