"""PDF Export - Convert DOCX to PDF"""

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from pathlib import Path
from datetime import datetime

class PDFExporter:
    """Export proposals to PDF"""
    
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._create_custom_styles()
    
    def _create_custom_styles(self):
        """Create custom styles for PDF"""
        # Title style
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=24,
            alignment=TA_CENTER,
            spaceAfter=30,
            textColor=colors.HexColor('#1a73e8')
        ))
        
        # Heading style
        self.styles.add(ParagraphStyle(
            name='CustomHeading',
            parent=self.styles['Heading2'],
            fontSize=16,
            spaceBefore=20,
            spaceAfter=10,
            textColor=colors.HexColor('#333333')
        ))
        
        # Normal text
        self.styles.add(ParagraphStyle(
            name='CustomNormal',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            leading=14
        ))
    
    def convert_docx_to_pdf(self, docx_path: str, pdf_path: str = None) -> str:
        """Convert DOCX to PDF"""
        
        docx_path = Path(docx_path)
        
        if not docx_path.exists():
            raise FileNotFoundError(f"DOCX not found: {docx_path}")
        
        if not pdf_path:
            pdf_path = docx_path.parent / f"{docx_path.stem}.pdf"
        
        # Read DOCX content
        doc = Document(docx_path)
        
        # Extract text and structure
        content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Check if it's a heading
                if paragraph.style.name.startswith('Heading'):
                    if paragraph.style.name == 'Heading 1':
                        content.append(Paragraph(paragraph.text, self.styles['CustomTitle']))
                    else:
                        content.append(Paragraph(paragraph.text, self.styles['CustomHeading']))
                else:
                    content.append(Paragraph(paragraph.text, self.styles['CustomNormal']))
                
                content.append(Spacer(1, 0.1 * inch))
        
        # Extract tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.replace('\n', ' ')
                    row_data.append(Paragraph(cell_text, self.styles['CustomNormal']))
                table_data.append(row_data)
            
            if table_data:
                t = Table(table_data)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                content.append(t)
                content.append(Spacer(1, 0.2 * inch))
        
        # Generate PDF
        doc = SimpleDocTemplate(
            str(pdf_path),
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        doc.build(content)
        
        print(f"✅ PDF saved: {pdf_path}")
        return str(pdf_path)

# Quick test
if __name__ == "__main__":
    exporter = PDFExporter()
    # Test with existing proposal
    test_file = "outputs/proposals/proposal_*.docx"
    import glob
    files = glob.glob(test_file)
    if files:
        exporter.convert_docx_to_pdf(files[0])